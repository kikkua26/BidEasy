"""文档导入解析 API"""

import os
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from app.core.config import settings
from app.core.dependencies import get_db
from app.schemas import ResponseModel, DocumentResponse, DocumentParseResult
from app.db.models import Document, Project

router = APIRouter(prefix="/api/v1/projects/{project_id}/documents", tags=["文档管理"])


async def _get_project(project_id: str, db: AsyncSession) -> Project:
    """获取项目（公共校验）"""
    result = await db.execute(select(Project).where(Project.id == project_id))
    project = result.scalar_one_or_none()
    if not project:
        raise HTTPException(status_code=404, detail="项目不存在")
    return project


@router.get("", response_model=ResponseModel)
async def list_documents(
    project_id: str,
    db: AsyncSession = Depends(get_db),
):
    """获取项目文档列表"""
    await _get_project(project_id, db)
    result = await db.execute(
        select(Document).where(Document.project_id == project_id).order_by(Document.created_at.desc())
    )
    docs = result.scalars().all()
    return ResponseModel(data=[DocumentResponse.model_validate(d) for d in docs])


@router.post("/upload", response_model=ResponseModel)
async def upload_document(
    project_id: str,
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
):
    """上传招文文件（PDF/Word/TXT）"""
    await _get_project(project_id, db)

    # 校验文件类型
    filename = file.filename or "unknown"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in ("pdf", "docx", "doc", "txt"):
        raise HTTPException(status_code=400, detail="仅支持 PDF / Word / TXT 格式")

    # 校验文件大小
    content = await file.read()
    max_size = settings.MAX_UPLOAD_SIZE_MB * 1024 * 1024
    if len(content) > max_size:
        raise HTTPException(status_code=400, detail=f"文件超过 {settings.MAX_UPLOAD_SIZE_MB}MB 限制")

    # 保存文件
    upload_dir = os.path.join(settings.UPLOAD_DIR, project_id)
    os.makedirs(upload_dir, exist_ok=True)
    file_path = os.path.join(upload_dir, filename)

    with open(file_path, "wb") as f:
        f.write(content)

    # 提取文本
    raw_text = ""
    try:
        if ext == "pdf":
            raw_text = await _parse_pdf(file_path)
        elif ext in ("docx", "doc"):
            raw_text = await _parse_docx(file_path)
        else:
            raw_text = content.decode("utf-8", errors="ignore")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件解析失败: {str(e)}")

    # 存储记录
    doc = Document(
        project_id=project_id,
        file_name=filename,
        file_type=ext,
        raw_text=raw_text,
    )
    db.add(doc)
    await db.flush()

    return ResponseModel(
        data=DocumentParseResult(
            file_name=filename,
            file_type=ext,
            raw_text=raw_text[:5000],  # 预览前5000字
            word_count=len(raw_text),
            page_count=None,
        ).model_dump()
    )


@router.post("/input", response_model=ResponseModel)
async def input_document_text(
    project_id: str,
    body: dict,
    db: AsyncSession = Depends(get_db),
):
    """手动输入/粘贴招文文本"""
    await _get_project(project_id, db)
    text = body.get("text", "").strip()
    if not text:
        raise HTTPException(status_code=400, detail="文本内容不能为空")

    doc = Document(
        project_id=project_id,
        file_name="手动输入",
        file_type="text",
        raw_text=text,
    )
    db.add(doc)
    await db.flush()
    return ResponseModel(
        data=DocumentParseResult(
            file_name="手动输入",
            file_type="text",
            raw_text=text[:5000],
            word_count=len(text),
            page_count=None,
        ).model_dump()
    )


@router.delete("/{doc_id}", response_model=ResponseModel)
async def delete_document(
    project_id: str,
    doc_id: str,
    db: AsyncSession = Depends(get_db),
):
    """删除文档"""
    await _get_project(project_id, db)
    result = await db.execute(
        select(Document).where(Document.id == doc_id, Document.project_id == project_id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")
    await db.delete(doc)
    return ResponseModel(message="文档已删除")


# ── 文档解析 ──

async def _parse_pdf(file_path: str) -> str:
    """解析PDF文件"""
    import fitz  # PyMuPDF
    text_parts = []
    with fitz.open(file_path) as doc:
        for page in doc:
            text_parts.append(page.get_text())
    return "\n\n".join(text_parts)


async def _parse_docx(file_path: str) -> str:
    """解析Word文件"""
    from docx import Document as DocxDocument
    doc = DocxDocument(file_path)
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    return "\n".join(text_parts)
